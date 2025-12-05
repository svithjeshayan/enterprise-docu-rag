# app_improved.py — PHASE 3 IMPROVED: Production-Ready Multi-User RAG Chatbot
import os
os.chdir(os.path.dirname(os.path.abspath(__file__)))

import streamlit as st
from pathlib import Path
import hashlib
import sqlite3
import json
import uuid
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Tuple
import fitz  # PyMuPDF
import pytesseract
from pdf2image import convert_from_path
import docx  # === NEW: For Word Documents ===
import threading
from contextlib import contextmanager
import logging

# LangChain + Vector DB
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS

import streamlit_authenticator as stauth
import yaml
from yaml.loader import SafeLoader

# === FIX INVISIBLE SOURCES - INJECT CUSTOM CSS ===
st.markdown("""
<style>
    .source-box {
        padding: 12px 16px;
        margin: 8px 0;
        border-left: 4px solid #4CAF50;
        border-radius: 0 8px 8px 0;
        background-color: var(--text-color) !important;   /* this is the trick */
        background-color: rgba(0, 0, 0, 0.05) !important; /* fallback */
        color: var(--text-color) !important;
        font-size: 0.95rem;
        line-height: 1.5;
    }
    .source-box strong {
        color: #4CAF50;
    }

    /* Force correct colors in both themes */
    [data-testid="stExpander"] .source-box {
        background-color: var(--background-color) !important;
        color: var(--text-color) !important;
    }

    /* Very safe fallback for light theme */
    @media (prefers-color-scheme: light) {
        .source-box {
            background-color: #f8f9fa !important;
            color: #212529 !important;
            border-left-color: #28a745 !important;
        }
        .source-box strong { color: #28a745 !important; }
    }

    /* Very safe fallback for dark theme */
    @media (prefers-color-scheme: dark) {
        .source-box {
            background-color: #2d2d2d !important;
            color: #e0e0e0 !important;
            border-left-color: #4ade80 !important;
        }
        .source-box strong { color: #4ade80 !important; }
    }
</style>
""", unsafe_allow_html=True)

# === LOGGING SETUP ===
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('chatbot.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# === CONFIG ===
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"E:\Downloads\Release-25.11.0-0\poppler-25.11.0\Library\bin"
DOCS_FOLDER = Path(r"Z:\Chat bot testing")
DB_PATH = "chatbot.db"
INDEX_PATH = Path("faiss_index")

# Rate limiting config
RATE_LIMIT_QUERIES = 20  # queries per time window
RATE_LIMIT_WINDOW = 60  # seconds

# Embedding & LLM with error handling
try:
    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
except Exception as e:
    logger.error(f"Failed to initialize OpenAI clients: {e}")
    st.error("⚠️ Failed to initialize AI models. Please check your API keys in secrets.")
    st.stop()

# Thread lock for vector store operations
vs_lock = threading.Lock()

# === DATABASE SETUP ===
def init_db():
    """Initialize SQLite database with proper schema"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    # Users table with UUID and roles
    c.execute("""CREATE TABLE IF NOT EXISTS users (
                user_id TEXT PRIMARY KEY,
                username TEXT UNIQUE,
                name TEXT,
                email TEXT,
                role TEXT DEFAULT 'user',
                folders_access TEXT,
                created_at TEXT,
                last_login TEXT,
                status TEXT DEFAULT 'active'
    )""")
    
    # Chats table with user isolation
    c.execute("""CREATE TABLE IF NOT EXISTS chats (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT,
                timestamp TEXT,
                role TEXT,
                content TEXT,
                FOREIGN KEY (user_id) REFERENCES users(user_id)
    )""")
    
    # Documents table with content hash
    c.execute("""CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT UNIQUE,
                file_hash TEXT,
                content_hash TEXT,
                processed_at TEXT,
                chunk_count INTEGER,
                status TEXT DEFAULT 'active'
    )""")
    
    # Document chunks for tracking
    c.execute("""CREATE TABLE IF NOT EXISTS document_chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id INTEGER,
                chunk_index INTEGER,
                content_hash TEXT,
                FOREIGN KEY (doc_id) REFERENCES documents(id)
    )""")
    
    # Rate limiting table
    c.execute("""CREATE TABLE IF NOT EXISTS rate_limits (
                user_id TEXT,
                timestamp TEXT,
                PRIMARY KEY (user_id, timestamp)
    )""")
    
    # Audit log table
    c.execute("""CREATE TABLE IF NOT EXISTS audit_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT,
                action TEXT,
                details TEXT,
                timestamp TEXT
    )""")
    
    # Create indexes for performance
    c.execute("CREATE INDEX IF NOT EXISTS idx_chats_user ON chats(user_id, timestamp)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_rate_limits ON rate_limits(user_id, timestamp)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_users_status ON users(status)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_audit_log ON audit_log(user_id, timestamp)")
    
    conn.commit()
    conn.close()

init_db()

# USER PERMISSIONS AND ROLES
USER_PERMISSIONS = {
    "john":  ["public", "hr", "finance"],
    "sarah": ["public", "hr"],
    "guest": ["public"],
    "admin": ["public", "hr", "finance", "admin"]  # Admin has access to all
}

USER_ROLES = {
    "john": "user",
    "sarah": "user", 
    "guest": "user",
    "admin": "admin"  # Admin role
}

# === ADMIN FUNCTIONS ===
def log_audit(user_id: str, action: str, details: str):
    """Log admin actions to audit trail"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""INSERT INTO audit_log (user_id, action, details, timestamp) 
                VALUES (?, ?, ?, ?)""",
              (user_id, action, details, datetime.now().isoformat()))
    conn.commit()
    conn.close()

def add_user_to_db(username: str, name: str, email: str, role: str, folders: List[str]):
    """Add user to database"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    try:
        c.execute("""INSERT INTO users (user_id, username, name, email, role, folders_access, created_at, status)
                    VALUES (?, ?, ?, ?, ?, ?, ?, 'active')""",
                  (str(uuid.uuid4()), username, name, email, role, json.dumps(folders), datetime.now().isoformat()))
        conn.commit()
        log_audit(user_id, "USER_CREATED", f"Created user: {username}")
        return True
    except sqlite3.IntegrityError:
        return False
    finally:
        conn.close()

def update_user_in_db(username: str, name: str, email: str, role: str, folders: List[str]):
    """Update user in database"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""UPDATE users SET name=?, email=?, role=?, folders_access=? 
                WHERE username=?""",
              (name, email, role, json.dumps(folders), username))
    conn.commit()
    conn.close()
    log_audit(user_id, "USER_UPDATED", f"Updated user: {username}")

def deactivate_user_in_db(username: str):
    """Deactivate user (soft delete)"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("UPDATE users SET status='inactive' WHERE username=?", (username,))
    conn.commit()
    conn.close()
    log_audit(user_id, "USER_DEACTIVATED", f"Deactivated user: {username}")

def get_all_users():
    """Get all users from database"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""SELECT username, name, email, role, folders_access, status, created_at, last_login 
                FROM users WHERE status='active' ORDER BY created_at DESC""")
    users = c.fetchall()
    conn.close()
    return users

def sync_user_to_db(username: str, name: str):
    """Sync authenticated user to database"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT user_id FROM users WHERE username=?", (username,))
    exists = c.fetchone()
    
    if not exists:
        # Create user record
        folders = json.dumps(USER_PERMISSIONS.get(username, ["public"]))
        role = USER_ROLES.get(username, "user")
        c.execute("""INSERT INTO users (user_id, username, name, role, folders_access, created_at, last_login, status)
                    VALUES (?, ?, ?, ?, ?, ?, ?, 'active')""",
                  (str(uuid.uuid4()), username, name, role, folders, datetime.now().isoformat(), 
                   datetime.now().isoformat()))
    else:
        # Update last login
        c.execute("UPDATE users SET last_login=? WHERE username=?",
                  (datetime.now().isoformat(), username))
    
    conn.commit()
    conn.close()

def update_credentials_yaml(username: str, name: str, email: str, password: str):
    """Update credentials.yaml file with new/updated user"""
    with open('credentials.yaml', 'r') as file:
        config = yaml.load(file, Loader=SafeLoader)
    
    # Hash password
    hashed_pw = stauth.Hasher([password]).generate()[0]
    
    # Add or update user
    config['credentials']['usernames'][username] = {
        'email': email,
        'name': name,
        'password': hashed_pw
    }
    
    # Save back to file
    with open('credentials.yaml', 'w') as file:
        yaml.dump(config, file, default_flow_style=False)

def remove_user_from_yaml(username: str):
    """Remove user from credentials.yaml"""
    with open('credentials.yaml', 'r') as file:
        config = yaml.load(file, Loader=SafeLoader)
    
    if username in config['credentials']['usernames']:
        del config['credentials']['usernames'][username]
    
    with open('credentials.yaml', 'w') as file:
        yaml.dump(config, file, default_flow_style=False)

# === AZURE AD / LDAP INTEGRATION (OPTIONAL) ===
def authenticate_with_azure_ad(username: str, password: str) -> bool:
    """
    Authenticate user with Azure AD / Microsoft Entra ID
    Requires: pip install msal
    
    To enable:
    1. Register app in Azure Portal
    2. Get CLIENT_ID, TENANT_ID, CLIENT_SECRET
    3. Uncomment and configure below
    """
    # Uncomment to enable Azure AD authentication
    """
    try:
        from msal import ConfidentialClientApplication
        
        CLIENT_ID = os.getenv("AZURE_CLIENT_ID")
        TENANT_ID = os.getenv("AZURE_TENANT_ID")
        CLIENT_SECRET = os.getenv("AZURE_CLIENT_SECRET")
        
        authority = f"https://login.microsoftonline.com/{TENANT_ID}"
        app = ConfidentialClientApplication(
            CLIENT_ID,
            authority=authority,
            client_credential=CLIENT_SECRET
        )
        
        result = app.acquire_token_by_username_password(
            username=username,
            password=password,
            scopes=["User.Read"]
        )
        
        if "access_token" in result:
            logger.info(f"Azure AD auth successful for {username}")
            return True
        else:
            logger.warning(f"Azure AD auth failed for {username}")
            return False
            
    except Exception as e:
        logger.error(f"Azure AD authentication error: {e}")
        return False
    """
    return False  # Disabled by default

def authenticate_with_ldap(username: str, password: str) -> bool:
    """
    Authenticate user with LDAP (Active Directory)
    Requires: pip install ldap3
    
    To enable:
    1. Configure LDAP server details
    2. Uncomment and configure below
    """
    # Uncomment to enable LDAP authentication
    """
    try:
        from ldap3 import Server, Connection, ALL
        
        LDAP_SERVER = os.getenv("LDAP_SERVER", "ldap://your-domain.com")
        LDAP_BASE_DN = os.getenv("LDAP_BASE_DN", "DC=company,DC=com")
        
        server = Server(LDAP_SERVER, get_info=ALL)
        user_dn = f"CN={username},{LDAP_BASE_DN}"
        
        conn = Connection(server, user=user_dn, password=password)
        
        if conn.bind():
            logger.info(f"LDAP auth successful for {username}")
            conn.unbind()
            return True
        else:
            logger.warning(f"LDAP auth failed for {username}")
            return False
            
    except Exception as e:
        logger.error(f"LDAP authentication error: {e}")
        return False
    """
    return False  # Disabled by default
def check_rate_limit(user_id: str) -> Tuple[bool, int]:
    """Check if user has exceeded rate limit. Returns (allowed, remaining_queries)"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    cutoff_time = (datetime.now() - timedelta(seconds=RATE_LIMIT_WINDOW)).isoformat()
    
    # Clean old entries
    c.execute("DELETE FROM rate_limits WHERE timestamp < ?", (cutoff_time,))
    
    # Count recent queries
    c.execute("SELECT COUNT(*) FROM rate_limits WHERE user_id = ? AND timestamp >= ?", 
              (user_id, cutoff_time))
    count = c.fetchone()[0]
    
    if count >= RATE_LIMIT_QUERIES:
        conn.close()
        return False, 0
    
    # Log this query
    c.execute("INSERT INTO rate_limits (user_id, timestamp) VALUES (?, ?)",
              (user_id, datetime.now().isoformat()))
    conn.commit()
    conn.close()
    
    return True, RATE_LIMIT_QUERIES - count - 1

# === PHASE 4: SECURE AUTH + PERMISSIONS ===
# Load credentials
with open('credentials.yaml') as file:
    config = yaml.load(file, Loader=SafeLoader)

authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days']
)

# Login widget
authenticator.login(location='main')

# Check authentication status
if st.session_state.get("authentication_status") == False:
    st.error('Wrong username/password')
    st.stop()
if st.session_state.get("authentication_status") is None:
    st.warning('Please enter your credentials')
    st.stop()

# Get user info from session state
name = st.session_state.get("name")
username = st.session_state.get("username")
authentication_status = st.session_state.get("authentication_status")

# Sync user to database
sync_user_to_db(username, name)

authenticator.logout(location='sidebar')
st.sidebar.success(f"Welcome **{name}**")

allowed_folders = USER_PERMISSIONS.get(username, ["public"])
user_role = USER_ROLES.get(username, "user")
user_id = username  # Use username as ID
user_name = name

# === VECTOR STORE MANAGEMENT ===
@contextmanager
def vector_store_session():
    """Context manager for thread-safe vector store operations"""
    vs_lock.acquire()
    try:
        yield
    finally:
        vs_lock.release()

def load_or_create_vectorstore() -> FAISS:
    """Load existing FAISS index or create new one"""
    index_file = INDEX_PATH / "index.faiss"
    if index_file.exists():
        try:
            return FAISS.load_local(
                str(INDEX_PATH), 
                embeddings, 
                allow_dangerous_deserialization=True
            )
        except Exception as e:
            logger.error(f"Failed to load FAISS index: {e}")
            # Backup corrupted index
            import shutil
            backup_path = INDEX_PATH / f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            shutil.move(str(INDEX_PATH), str(backup_path))
            logger.info(f"Corrupted index backed up to {backup_path}")
    
    # Create new index
    dummy_doc = Document(page_content="Initialization document", metadata={"source": "system"})
    return FAISS.from_documents([dummy_doc], embeddings)

# Initialize vector store in session state for thread safety
if "vectorstore" not in st.session_state:
    with vector_store_session():
        st.session_state.vectorstore = load_or_create_vectorstore()

# === EXTRACTORS FOR DIFFERENT FILE TYPES ===
def extract_text_from_docx(path: Path) -> Optional[str]:
    """Extract text from .docx files"""
    try:
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"DOCX extraction failed for {path.name}: {e}")
        return None

def extract_text_from_txt(path: Path) -> Optional[str]:
    """Extract text from .txt files"""
    try:
        return path.read_text(encoding='utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"TXT extraction failed for {path.name}: {e}")
        return None

def extract_text_from_pdf(pdf_path: Path) -> Optional[str]:
    """Native PDF extraction with OCR fallback"""
    try:
        # Try native text extraction first
        doc = fitz.open(pdf_path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        
        # If sufficient text found, return it
        if len(text.strip()) > 100:
            logger.info(f"Extracted {len(text)} chars from {pdf_path.name} (native)")
            return text
    except Exception as e:
        logger.warning(f"Native extraction failed for {pdf_path.name}: {e}")
    
    # Fallback to OCR
    try:
        logger.info(f"Starting OCR for {pdf_path.name}")
        images = convert_from_path(
            str(pdf_path), 
            dpi=300, 
            poppler_path=POPPLER_PATH,
        )
        
        ocr_text = []
        for i, img in enumerate(images):
            try:
                page_text = pytesseract.image_to_string(img, lang='eng')
                ocr_text.append(page_text)
                logger.info(f"OCR page {i+1}/{len(images)} complete")
            except Exception as e:
                logger.error(f"OCR failed on page {i+1}: {e}")
                continue
        
        full_text = "\n\n".join(ocr_text)
        logger.info(f"OCR complete: {len(full_text)} chars from {pdf_path.name}")
        return full_text if full_text.strip() else None
        
    except Exception as e:
        logger.error(f"OCR failed for {pdf_path.name}: {e}")
        return None

# === MAIN DISPATCHER ===
def extract_text(file_path: Path) -> Optional[str]:
    ext = file_path.suffix.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        logger.warning(f"Unsupported file type: {file_path.name}")
        return None

def get_content_hash(text: str) -> str:
    """Generate hash of document content (not file)"""
    return hashlib.sha256(text.encode()).hexdigest()

def get_file_hash(path: Path) -> str:
    """Generate hash of file bytes"""
    return hashlib.md5(path.read_bytes()).hexdigest()

# === DOCUMENT PROCESSING ===
def rebuild_vectorstore():
    """Rebuild/update vector store with new or modified documents"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    with st.status("🔄 Updating knowledge base...") as status:
        docs_to_add = []
        docs_to_remove = []
        
        # Check for removed documents
        c.execute("SELECT id, filename FROM documents WHERE status='active'")
        existing_docs = {row[1]: row[0] for row in c.fetchall()}
        
        # === FIXED: Now allowed_folders is defined ===
        current_files_paths = []
        for folder_name in allowed_folders:
            folder_path = DOCS_FOLDER / folder_name
            if not folder_path.exists():
                logger.warning(f"Folder not found: {folder_path}")
                continue
            
            # Get all supported file types from this folder
            for ext in ['*.pdf', '*.docx', '*.txt']:
                current_files_paths.extend(folder_path.glob(ext))
        
        current_filenames = {f.name for f in current_files_paths}
        removed_files = set(existing_docs.keys()) - current_filenames
        
        if removed_files:
            status.write(f"📤 Removing {len(removed_files)} deleted documents...")
            for filename in removed_files:
                c.execute("UPDATE documents SET status='deleted' WHERE filename=?", (filename,))
                logger.info(f"Marked {filename} as deleted")
        
        # Process new/updated documents
        processed = 0
        failed = 0
        
        for file_path in current_files_paths:
            file_hash = get_file_hash(file_path)
            
            c.execute("SELECT file_hash, content_hash FROM documents WHERE filename=?", (file_path.name,))
            row = c.fetchone()
            
            # Skip if unchanged
            if row and row[0] == file_hash:
                continue
            
            status.write(f"📄 Processing: {file_path.name}")
            
            try:
                # Extract text using the dispatcher
                text = extract_text(file_path)
                if not text:
                    logger.error(f"No text extracted from {file_path.name}")
                    failed += 1
                    continue
                
                content_hash = get_content_hash(text)
                
                # Skip if content unchanged (file metadata changed but not content)
                if row and row[1] == content_hash:
                    c.execute("UPDATE documents SET file_hash=? WHERE filename=?", 
                             (file_hash, file_path.name))
                    continue
                
                # Split into chunks
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=800,
                    chunk_overlap=150,
                    length_function=len,
                    separators=["\n\n", "\n", ". ", " ", ""]
                )
                chunks = splitter.split_text(text)
                
                # Create documents with rich metadata
                for i, chunk in enumerate(chunks):
                    docs_to_add.append(Document(
                        page_content=chunk,
                        metadata={
                            "source": file_path.name,
                            "chunk": i + 1,
                            "total_chunks": len(chunks),
                            "updated": datetime.now().isoformat(),
                            "content_hash": get_content_hash(chunk)
                        }
                    ))
                
                # Update database
                c.execute("""INSERT OR REPLACE INTO documents 
                            (filename, file_hash, content_hash, processed_at, chunk_count, status)
                            VALUES (?, ?, ?, ?, ?, 'active')""",
                         (file_path.name, file_hash, content_hash, datetime.now().isoformat(), len(chunks)))
                
                processed += 1
                logger.info(f"Processed {file_path.name}: {len(chunks)} chunks")
                
            except Exception as e:
                logger.error(f"Failed to process {file_path.name}: {e}")
                failed += 1
                continue
        
        # Update vector store
        if docs_to_add:
            status.write(f"🔧 Adding {len(docs_to_add)} chunks to vector store...")
            
            with vector_store_session():
                st.session_state.vectorstore.add_documents(docs_to_add)
                st.session_state.vectorstore.save_local(str(INDEX_PATH))
            
            status.update(label=f"✅ Knowledge base updated! ({processed} docs processed)", 
                          state="complete")
            st.success(f"✅ Processed {processed} documents, {len(docs_to_add)} chunks added")
            if failed > 0:
                st.warning(f"⚠️ {failed} documents failed to process")
        else:
            status.update(label="✅ All documents up to date", state="complete")
            st.info("ℹ️ No new documents to process")
    
    conn.commit()
    conn.close()

# === CHAT HISTORY MANAGEMENT ===
def load_chat_history(limit: int = 50) -> List[Dict[str, str]]:
    """Load recent chat history for current user"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""SELECT role, content FROM chats 
                WHERE user_id=? 
                ORDER BY timestamp DESC 
                LIMIT ?""", (user_id, limit))
    messages = [{"role": r[0], "content": r[1]} for r in reversed(c.fetchall())]
    conn.close()
    return messages

def save_message(role: str, content: str):
    """Save message to database"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""INSERT INTO chats (user_id, timestamp, role, content) 
                VALUES (?, ?, ?, ?)""",
              (user_id, datetime.now().isoformat(), role, content))
    conn.commit()
    conn.close()

def get_conversation_context(messages: List[Dict], max_turns: int = 5) -> str:
    """Build conversation context from recent messages"""
    if not messages:
        return ""
    
    recent = messages[-max_turns*2:] if len(messages) > max_turns*2 else messages
    context_lines = []
    
    for msg in recent:
        role = "User" if msg["role"] == "user" else "Assistant"
        context_lines.append(f"{role}: {msg['content']}")
    
    return "\n".join(context_lines)

# Initialize chat history
if "messages" not in st.session_state:
    st.session_state.messages = load_chat_history()

# === MAIN UI ===
st.set_page_config(
    page_title="Enterprise RAG Chatbot",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .stChatMessage { padding: 1rem; }
    .source-box { 
        background: #f0f2f6; 
        padding: 0.5rem; 
        border-radius: 0.5rem; 
        margin: 0.5rem 0;
        font-size: 0.9rem;
    }
</style>
""", unsafe_allow_html=True)

st.title("Secure Enterprise Document Assistant")
st.caption(f"Logged in as **{name}** | Access: {', '.join(allowed_folders)}")

# === SIDEBAR ===
with st.sidebar:
    st.header("⚙️ Controls")
    
    # === ADMIN PANEL ===
    if user_role == "admin":
        with st.expander("👑 Admin Panel", expanded=False):
            st.subheader("User Management")
            
            admin_action = st.radio(
                "Action",
                ["View Users", "Add User", "Edit User", "Deactivate User", "View Audit Log"],
                key="admin_action"
            )
            
            if admin_action == "View Users":
                st.markdown("#### Active Users")
                users = get_all_users()
                if users:
                    for user in users:
                        username_db, name_db, email_db, role_db, folders_db, status_db, created_db, last_login_db = user
                        folders_list = json.loads(folders_db) if folders_db else []
                        
                        st.markdown(f"""
                        **{name_db}** (@{username_db})  
                        📧 {email_db or 'N/A'} | 🔑 {role_db}  
                        📁 Access: {', '.join(folders_list)}  
                        🕐 Created: {created_db[:10] if created_db else 'N/A'}
                        """)
                        st.divider()
                else:
                    st.info("No users found")
            
            elif admin_action == "Add User":
                st.markdown("#### Create New User")
                with st.form("add_user_form"):
                    new_username = st.text_input("Username*")
                    new_name = st.text_input("Full Name*")
                    new_email = st.text_input("Email*")
                    new_password = st.text_input("Password*", type="password")
                    new_role = st.selectbox("Role", ["user", "admin"])
                    new_folders = st.multiselect(
                        "Folder Access",
                        ["public", "hr", "finance", "admin"],
                        default=["public"]
                    )
                    
                    if st.form_submit_button("Create User"):
                        if new_username and new_name and new_email and new_password:
                            # Add to database
                            if add_user_to_db(new_username, new_name, new_email, new_role, new_folders):
                                # Add to credentials.yaml
                                update_credentials_yaml(new_username, new_name, new_email, new_password)
                                st.success(f"✅ User {new_username} created successfully!")
                                st.info("⚠️ User must log out and back in for changes to take effect")
                            else:
                                st.error("❌ Username already exists")
                        else:
                            st.error("❌ All fields are required")
            
            elif admin_action == "Edit User":
                st.markdown("#### Edit Existing User")
                users = get_all_users()
                usernames = [u[0] for u in users]
                
                if usernames:
                    selected_user = st.selectbox("Select User", usernames)
                    
                    # Get current user data
                    conn = sqlite3.connect(DB_PATH)
                    c = conn.cursor()
                    c.execute("""SELECT name, email, role, folders_access 
                                FROM users WHERE username=?""", (selected_user,))
                    user_data = c.fetchone()
                    conn.close()
                    
                    if user_data:
                        current_name, current_email, current_role, current_folders = user_data
                        current_folders_list = json.loads(current_folders) if current_folders else ["public"]
                        
                        with st.form("edit_user_form"):
                            edit_name = st.text_input("Full Name", value=current_name)
                            edit_email = st.text_input("Email", value=current_email or "")
                            edit_role = st.selectbox("Role", ["user", "admin"], 
                                                    index=0 if current_role == "user" else 1)
                            edit_folders = st.multiselect(
                                "Folder Access",
                                ["public", "hr", "finance", "admin"],
                                default=current_folders_list
                            )
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                if st.form_submit_button("Update User"):
                                    update_user_in_db(selected_user, edit_name, edit_email, edit_role, edit_folders)
                                    st.success(f"✅ User {selected_user} updated!")
                            
                            with col2:
                                new_pw = st.text_input("New Password (optional)", type="password")
                                if st.form_submit_button("Reset Password") and new_pw:
                                    update_credentials_yaml(selected_user, edit_name, edit_email, new_pw)
                                    st.success("✅ Password reset successfully!")
                else:
                    st.info("No users to edit")
            
            elif admin_action == "Deactivate User":
                st.markdown("#### Deactivate User")
                users = get_all_users()
                usernames = [u[0] for u in users if u[0] != username]  # Can't deactivate self
                
                if usernames:
                    deactivate_user = st.selectbox("Select User", usernames)
                    
                    if st.button("⚠️ Deactivate User", type="secondary"):
                        deactivate_user_in_db(deactivate_user)
                        remove_user_from_yaml(deactivate_user)
                        st.success(f"✅ User {deactivate_user} deactivated")
                        st.info("⚠️ Refresh the page to see changes")
                else:
                    st.info("No users to deactivate")
            
            elif admin_action == "View Audit Log":
                st.markdown("#### Audit Log (Last 50 Actions)")
                conn = sqlite3.connect(DB_PATH)
                c = conn.cursor()
                c.execute("""SELECT user_id, action, details, timestamp 
                            FROM audit_log ORDER BY timestamp DESC LIMIT 50""")
                logs = c.fetchall()
                conn.close()
                
                if logs:
                    for log in logs:
                        user_log, action_log, details_log, time_log = log
                        st.text(f"[{time_log[:19]}] {user_log}: {action_log} - {details_log}")
                else:
                    st.info("No audit logs yet")
    
    st.markdown("---")
    
    if st.button("🔄 Rebuild Knowledge Base", use_container_width=True):
        rebuild_vectorstore()
        st.rerun()
    
    if st.button("🗑️ Clear My Chat History", use_container_width=True):
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("DELETE FROM chats WHERE user_id=?", (user_id,))
        conn.commit()
        conn.close()
        st.session_state.messages = []
        st.success("Chat history cleared!")
        st.rerun()
    
    st.markdown("---")
    st.header("📊 Statistics")
    
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    c.execute("SELECT COUNT(*) FROM documents WHERE status='active'")
    doc_count = c.fetchone()[0]
    
    c.execute("SELECT SUM(chunk_count) FROM documents WHERE status='active'")
    chunk_count = c.fetchone()[0] or 0
    
    c.execute("SELECT COUNT(*) FROM chats WHERE user_id=?", (user_id,))
    msg_count = c.fetchone()[0]
    
    conn.close()
    
    st.metric("📄 Active Documents", doc_count)
    st.metric("🧩 Total Chunks", chunk_count)
    st.metric("💬 Your Messages", msg_count)
    
    # Rate limit info
    allowed_queries, remaining = check_rate_limit(user_id + "_check")
    st.metric("⏱️ Queries Remaining", f"{remaining}/{RATE_LIMIT_QUERIES}")

# Auto-rebuild on first load
if not (INDEX_PATH / "index.faiss").exists():
    with st.spinner("Building initial knowledge base..."):
        rebuild_vectorstore()

# === CHAT DISPLAY ===
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# === CHAT INPUT ===
if prompt := st.chat_input("💬 Ask about your documents...", key="chat_input"):
    # Check rate limit
    allowed_queries, remaining = check_rate_limit(user_id)
    if not allowed_queries:
        st.error(f"⚠️ Rate limit exceeded. Please wait {RATE_LIMIT_WINDOW} seconds before sending more messages.")
        st.stop()
    
    # Display user message
    st.session_state.messages.append({"role": "user", "content": prompt})
    save_message("user", prompt)
    with st.chat_message("user"):
        st.markdown(prompt)

    # Generate response
    with st.chat_message("assistant"):
        with st.spinner("🔍 Searching documents..."):
            try:
                # Perform semantic search
                with vector_store_session():
                    results = st.session_state.vectorstore.similarity_search_with_score(
                        prompt, 
                        k=8
                    )
                
                # Filter by relevance score (lower is better for FAISS)
                relevant_results = [(doc, score) for doc, score in results if score < 1.0]
                
                if not relevant_results:
                    answer = "I couldn't find relevant information in the current documents to answer your question."
                    st.markdown(answer)
                else:
                    # Build context
                    context_parts = []
                    for doc, score in relevant_results:
                        source = doc.metadata.get('source', 'Unknown')
                        chunk_num = doc.metadata.get('chunk', '?')
                        context_parts.append(
                            f"[Source: {source}, Chunk {chunk_num}, Relevance: {1-score:.2f}]\n{doc.page_content}"
                        )
                    
                    context = "\n\n---\n\n".join(context_parts)
                    conversation_history = get_conversation_context(st.session_state.messages[:-1])
                    
                    # Build prompt with conversation context
                    system_prompt = f"""You are an expert document assistant. Answer questions using ONLY the provided context.

CONVERSATION HISTORY:
{conversation_history if conversation_history else "No previous conversation"}

DOCUMENT CONTEXT:
{context}

INSTRUCTIONS:
- Provide accurate, professional answers based solely on the context
- If information is not in the context, say: "I don't have this information in the current documents"
- Cite sources when possible (e.g., "According to [document name]...")
- Keep answers concise but complete
- Maintain conversation continuity when relevant

QUESTION: {prompt}"""
                    
                    # Get LLM response
                    response = llm.invoke(system_prompt)
                    answer = response.content
                    
                    st.markdown(answer)
                    
                    # Show sources
                    with st.expander("View Sources", expanded=False):
                        for i, (doc, score) in enumerate(relevant_results, 1):
                            source = doc.metadata.get('source', 'Unknown document')
                            chunk = doc.metadata.get('chunk', '?')
                            preview = doc.page_content.replace("\n", " ").strip()
                            
                            st.markdown(f"""
                            <div class="source-box">
                                <strong>#{i} — {source}</strong> (Chunk {chunk})<br>
                                <small>Relevance: <b>{1-score:.3f}</b></small><br>
                                {preview[:450]}{"..." if len(preview) > 450 else ""}
                            </div>
                            """, unsafe_allow_html=True)
                
                # Save assistant message
                st.session_state.messages.append({"role": "assistant", "content": answer})
                save_message("assistant", answer)
                
            except Exception as e:
                logger.error(f"Error generating response: {e}")
                error_msg = "I encountered an error while processing your question. Please try again."
                st.error(error_msg)
                st.session_state.messages.append({"role": "assistant", "content": error_msg})
                save_message("assistant", error_msg)

# Footer
st.markdown("---")
st.caption(f"🔒 Session ID: `{user_id[:8]}...` | Queries remaining: {remaining}/{RATE_LIMIT_QUERIES}")